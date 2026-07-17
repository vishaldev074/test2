import streamlit as st
import io
import os
import tempfile
import numpy as np
from PIL import Image, ImageDraw, ImageFont, ImageEnhance
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Pt, RGBColor
from moviepy.editor import VideoFileClip, ImageClip, CompositeVideoClip

# --- Helper Logic: Position Maps ---

def get_position_coords(base_w, base_h, wm_w, wm_h, position_name, padding=20):
    """Calculates top-left coordinates for watermarks based on text/logo boundaries."""
    if position_name == "Top Left":
        return padding, padding
    elif position_name == "Top Right":
        return base_w - wm_w - padding, padding
    elif position_name == "Bottom Left":
        return padding, base_h - wm_h - padding
    elif position_name == "Bottom Right":
        return base_w - wm_w - padding, base_h - wm_h - padding
    else:  # Center
        return (base_w - wm_w) // 2, (base_h - wm_h) // 2

# --- Core Processor Engine: Images ---

def apply_image_watermark(orig_img, wm_type, text, logo_file, pos, opacity, rotation, color, size):
    base = orig_img.convert("RGBA")
    wm_layer = Image.new("RGBA", base.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(wm_layer)
    
    if wm_type == "Text" and text:
        # Load font with the requested dynamic size scale
        try:
            font = ImageFont.load_default(size=size)
        except TypeError:
            font = ImageFont.load_default()
            
        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        
        # Translate the Hex color to an RGB layout array
        hex_c = color.lstrip('#')
        r, g, b = tuple(int(hex_c[i:i+2], 16) for i in (0, 2, 4))
        
        # Create separate canvas for text rotation and color fill
        text_canvas = Image.new("RGBA", (tw + 10, th + 10), (0, 0, 0, 0))
        text_draw = ImageDraw.Draw(text_canvas)
        text_draw.text((5, 5), text, font=font, fill=(r, g, b, int(opacity * 255)))
        rotated_txt = text_canvas.rotate(rotation, expand=True)
        
        x, y = get_position_coords(base.size[0], base.size[1], rotated_txt.size[0], rotated_txt.size[1], pos)
        wm_layer.paste(rotated_txt, (x, y), rotated_txt)
        
    elif wm_type == "Logo" and logo_file:
        logo = Image.open(logo_file).convert("RGBA")
        
        # Scale logo relative to the base image and size slider configuration
        scale_factor = size / 200.0
        new_width = max(10, int(base.size[0] * scale_factor))
        new_height = max(10, int(new_width * (logo.height / logo.width)))
        logo = logo.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Opacity overlay adjustments
        alpha = logo.split()[3]
        alpha = ImageEnhance.Brightness(alpha).enhance(opacity)
        logo.putalpha(alpha)
        
        rotated_logo = logo.rotate(rotation, expand=True)
        x, y = get_position_coords(base.size[0], base.size[1], rotated_logo.size[0], rotated_logo.size[1], pos)
        wm_layer.paste(rotated_logo, (x, y), rotated_logo)
        
    return Image.alpha_composite(base, wm_layer).convert("RGB")

# --- Core Processor Engine: PDFs ---

def create_pdf_overlay(wm_type, text, logo_bytes, pos, opacity, rotation, color, size):
    """Generates a temporary vector layer page using ReportLab canvas arrays."""
    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    can.saveState()
    
    # Establish Alpha/Opacity Layer settings
    can.setFillAlpha(opacity)
    
    # Basic translation coordinates map for standard ReportLab Letter Layout
    w, h = letter
    x, y = w / 2, h / 2
    if pos == "Top Left": x, y = 70, h - 70
    elif pos == "Top Right": x, y = w - 150, h - 70
    elif pos == "Bottom Left": x, y = 70, 70
    elif pos == "Bottom Right": x, y = w - 150, 70
    
    can.translate(x, y)
    can.rotate(rotation)
    
    if wm_type == "Text" and text:
        hex_c = color.lstrip('#')
        r, g, b = [int(hex_c[i:i+2], 16)/255.0 for i in (0, 2, 4)]
        can.setFillColorRGB(r, g, b)
        can.setFont("Helvetica-Bold", size)
        can.drawCentredString(0, 0, text)
        
    elif wm_type == "Logo" and logo_bytes:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tf:
            tf.write(logo_bytes)
            tf_name = tf.name
            
        logo_w, logo_h = size * 3, size * 3
        can.drawImage(tf_name, -logo_w/2, -logo_h/2, width=logo_w, height=logo_h, mask='auto')
        os.remove(tf_name)
        
    can.restoreState()
    can.save()
    packet.seek(0)
    return PdfReader(packet).pages[0]

# --- Core Processor Engine: Word Documents ---

def watermark_docx(doc_bytes, text, color, size):
    """Injects warning strings structurally into the document body headers."""
    doc = Document(io.BytesIO(doc_bytes))
    hex_c = color.lstrip('#') if color else "808080"
    
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""  # Clear any existing paragraph strings
        
        run = hp.add_run(f"[{text}] - SECURED WATERMARK")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16))
        hp.alignment = 2  # Right-aligned header context layout
    
    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()

# --- Application Frontend Interface ---

st.set_page_config(page_title="AI Watermark Studio", page_icon="🛡️", layout="wide")
st.title("🛡️ AI Watermark Studio")
st.markdown("Secure and verify your digital assets. Apply watermarks seamlessly to images, videos, and documentation structures.")

# --- Sidebar Parameter Controls ---
st.sidebar.title("⚙️ Watermark Configurations")
wm_type = st.sidebar.radio("Watermark Medium Style", ["Text", "Logo"])
wm_text = st.sidebar.text_input("Watermark Plain Text", value="DRAFT") if wm_type == "Text" else None
logo_data = st.sidebar.file_uploader("Upload Image Logo Layer", type=["png", "jpg", "jpeg"]) if wm_type == "Logo" else None

# New Color and Size Controllers
wm_color = st.sidebar.color_picker("Watermark Color", "#FFFFFF") if wm_type == "Text" else None
wm_size = st.sidebar.slider("Watermark Scale / Font Size", 10, 200, 50)

wm_pos = st.sidebar.selectbox("Position Coordinates Layout", ["Center", "Top Left", "Top Right", "Bottom Left", "Bottom Right"])
# Updated Opacity logic (0.01 step interval for smooth tracking)
wm_opacity = st.sidebar.slider("Opacity Alpha Scale", 0.0, 1.0, 0.8, 0.01)
wm_rot = st.sidebar.slider("Rotation Matrix Angle (Degrees)", -180, 180, 0, 5)

tab_img, tab_vid, tab_pdf, tab_docx = st.tabs(["🖼️ Images", "🎬 Videos", "📄 PDFs", "📝 Word Docs"])

# --- Tab 1: Image Processing Logic ---
with tab_img:
    st.header("Image Watermark Deck")
    img_uploads = st.file_uploader("Upload Target Images", type=["png", "jpg", "jpeg", "webp"], accept_multiple_files=True)
    if img_uploads:
        for item in img_uploads:
            pil_img = Image.open(item)
            res = apply_image_watermark(pil_img, wm_type, wm_text, logo_data, wm_pos, wm_opacity, wm_rot, wm_color, wm_size)
            
            st.image(res, caption=f"Preview: {item.name}", width=350)
            out_buf = io.BytesIO()
            res.save(out_buf, format="JPEG")
            st.download_button(f"Download Protected {item.name}", out_buf.getvalue(), file_name=f"wm_{item.name}", mime="image/jpeg")

# --- Tab 2: Video Processing Logic ---
with tab_vid:
    st.header("Video Composition Processing Deck")
    vid_file = st.file_uploader("Upload Target Video Asset", type=["mp4", "mov", "avi"])
    if vid_file:
        st.warning("Video composition tasks process across background threads; larger container tasks scale processing times.")
        if st.button("Apply Digital Video Stamp"):
            with st.spinner("Compiling and scaling frame layers..."):
                t_input = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4")
                t_input.write(vid_file.read())
                
                try:
                    clip = VideoFileClip(t_input.name)
                    dummy = Image.new("RGBA", (clip.size[0], clip.size[1]), (0,0,0,0))
                    wm_frame = apply_image_watermark(dummy, wm_type, wm_text, logo_data, wm_pos, wm_opacity, wm_rot, wm_color, wm_size)
                    
                    t_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                    wm_frame.save(t_img.name)
                    
                    wm_clip = ImageClip(t_img.name).set_duration(clip.duration).set_opacity(wm_opacity)
                    final_clip = CompositeVideoClip([clip, wm_clip.set_position("center" if wm_pos == "Center" else "left")])
                    
                    t_out = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4")
                    final_clip.write_videofile(t_out.name, fps=24, codec="libx264", audio_codec="aac", preset="ultrafast", logger=None)
                    
                    with open(t_out.name, "rb") as f:
                        st.video(f.read())
                        st.download_button("Download Secured MP4 Video", f.read(), file_name="secured_video.mp4", mime="video/mp4")
                        
                    clip.close()
                    final_clip.close()
                    os.remove(t_input.name)
                    os.remove(t_img.name)
                    os.remove(t_out.name)
                except Exception as e:
                    st.error(f"Video compositor error tracking: {e}")

# --- Tab 3: PDF Processing Logic ---
with tab_pdf:
    st.header("PDF Document Vector Injection Deck")
    pdf_file = st.file_uploader("Upload Target PDF Document", type=["pdf"])
    if pdf_file:
        if st.button("Inject Structural PDF Watermark"):
            try:
                reader = PdfReader(pdf_file)
                writer = PdfWriter()
                
                logo_bytes_data = logo_data.read() if logo_data else None
                overlay_page = create_pdf_overlay(wm_type, wm_text, logo_bytes_data, wm_pos, wm_opacity, wm_rot, wm_color, wm_size)
                
                for idx in range(len(reader.pages)):
                    page = reader.pages[idx]
                    page.merge_page(overlay_page)
                    writer.add_page(page)
                    
                pdf_out = io.BytesIO()
                writer.write(pdf_out)
                st.success("Vector layers injected successfully into document structure.")
                st.download_button("Download Processed PDF Asset", pdf_out.getvalue(), file_name="secured_document.pdf", mime="application/pdf")
            except Exception as e:
                st.error(f"PDF structural compiler failure loop notice: {e}")

# --- Tab 4: DOCX Processing Logic ---
with tab_docx:
    st.header("Word Document Global Track Header Injection Deck")
    docx_file = st.file_uploader("Upload Microsoft Word Document File", type=["docx"])
    if docx_file:
        if st.button("Apply Document Header Stamp"):
            try:
                stamp_text = wm_text if wm_text else "SECURED DATA"
                docx_res = watermark_docx(docx_file.read(), stamp_text, wm_color, wm_size)
                st.success("Global headers injected throughout XML document layers.")
                st.download_button("Download Processed DOCX File", docx_res, file_name="secured_document.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"XML structural alignment exception: {e}")
